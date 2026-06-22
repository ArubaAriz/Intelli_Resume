"""
Resume parser: PDF and DOCX → structured JSON.
Extracts contact info, education, skills, experience, projects, certifications.
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional

import fitz  # PyMuPDF
from docx import Document

from nlp_utils import extract_skills_from_text, NOISE_WORDS


# ── Text Extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(path: str) -> tuple[str, Dict]:
    """Extract text + ATS structural metadata from PDF."""
    doc = fitz.open(path)
    full_text = ""
    ats_flags = {
        "has_tables": False,
        "has_images": False,
        "multi_column": False,
        "page_count": doc.page_count,
        "font_count": 0,
    }
    fonts = set()
    
    for page in doc:
        # Check for images
        if page.get_images():
            ats_flags["has_images"] = True
        
        # Check for tables (simplified: look for table-like structures)
        blocks = page.get_text("blocks")
        
        # Multi-column detection: check x-positions of text blocks
        x_positions = [b[0] for b in blocks if b[4].strip()]
        if x_positions:
            x_range = max(x_positions) - min(x_positions)
            page_width = page.rect.width
            if x_range > page_width * 0.4:
                ats_flags["multi_column"] = True
        
        full_text += page.get_text() + "\n"
        
        # Font diversity
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        fonts.add(span.get("font", ""))
    
    ats_flags["font_count"] = len(fonts)
    doc.close()
    return full_text, ats_flags


def extract_text_from_docx(path: str) -> tuple[str, Dict]:
    """Extract text from DOCX."""
    doc = Document(path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    ats_flags = {
        "has_tables": len(doc.tables) > 0,
        "has_images": False,
        "multi_column": False,
        "page_count": 1,
        "font_count": 1,
    }
    return full_text, ats_flags


def extract_text(file_path: str) -> tuple[str, Dict]:
    path = Path(file_path)
    if path.suffix.lower() == ".pdf":
        return extract_text_from_pdf(file_path)
    elif path.suffix.lower() in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {path.suffix}")


# ── Contact Info ──────────────────────────────────────────────────────────────

def extract_contact_info(text: str) -> Dict:
    """Extract name, email, phone, LinkedIn, GitHub."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    
    email = None
    phone = None
    linkedin = None
    github = None
    name = None
    
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    phone_pattern = re.compile(r'[\+]?[1-9][\d\s\-\(\)]{8,14}')
    linkedin_pattern = re.compile(r'(?:linkedin\.com/in/|linkedin:\s*)([A-Za-z0-9\-]+)', re.I)
    github_pattern = re.compile(r'(?:github\.com/|github:\s*)([A-Za-z0-9\-]+)', re.I)
    
    for line in lines[:30]:  # contact info usually at top
        if not email:
            m = email_pattern.search(line)
            if m:
                email = m.group(0)
        
        if not phone:
            m = phone_pattern.search(line)
            if m:
                raw = m.group(0).strip()
                if len(re.sub(r'\D', '', raw)) >= 7:
                    phone = raw
        
        if not linkedin:
            m = linkedin_pattern.search(line)
            if m:
                linkedin = f"linkedin.com/in/{m.group(1)}"
        
        if not github:
            m = github_pattern.search(line)
            if m:
                github = f"github.com/{m.group(1)}"
    
    # Name: first non-empty, non-label line at top that looks like a name
    name_skip = {"resume", "curriculum vitae", "cv", "email:", "phone:", "contact"}
    for line in lines[:8]:
        l = line.strip()
        if (l.lower() not in name_skip 
                and not email_pattern.search(l)
                and not phone_pattern.search(l)
                and 2 <= len(l.split()) <= 5
                and not re.search(r'[|•|@]', l)
                and l[0].isupper()):
            name = l
            break
    
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
    }


# ── Section Detection ─────────────────────────────────────────────────────────

SECTION_HEADERS = {
    "education": ["education", "academic background", "qualifications", "academic history"],
    "experience": ["experience", "work experience", "employment", "work history", "professional experience", "internship"],
    "skills": ["skills", "technical skills", "core competencies", "technologies", "expertise"],
    "projects": ["projects", "personal projects", "key projects", "academic projects", "side projects"],
    "certifications": ["certifications", "certificates", "credentials", "courses", "training"],
    "achievements": ["achievements", "awards", "honors", "accomplishments", "recognition"],
    "summary": ["summary", "profile", "objective", "about", "professional summary"],
}

def split_into_sections(text: str) -> Dict[str, str]:
    """Split resume text into named sections."""
    lines = text.split('\n')
    sections: Dict[str, List[str]] = {k: [] for k in SECTION_HEADERS}
    sections["other"] = []
    current_section = "other"
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        line_lower = stripped.lower().rstrip(':').rstrip()
        matched_section = None
        for section, headers in SECTION_HEADERS.items():
            if line_lower in headers or any(line_lower.startswith(h) for h in headers):
                matched_section = section
                break
        
        if matched_section:
            current_section = matched_section
        else:
            sections[current_section].append(stripped)
    
    return {k: '\n'.join(v) for k, v in sections.items()}


# ── Education Parsing ─────────────────────────────────────────────────────────

def parse_education(edu_text: str) -> List[Dict]:
    degrees = []
    degree_patterns = [
        r'(B\.?Tech|Bachelor|B\.?E|BS|B\.?Sc)\.?\s+(?:of|in|-)?\s*([\w\s,&]+)',
        r'(M\.?Tech|Master|M\.?E|MS|M\.?Sc|MBA)\.?\s+(?:of|in|-)?\s*([\w\s,&]+)',
        r'(Ph\.?D|Doctorate)\.?\s+(?:in|-)?\s*([\w\s,&]+)',
    ]
    
    lines = edu_text.split('\n')
    for line in lines:
        if not line.strip():
            continue
        for pattern in degree_patterns:
            m = re.search(pattern, line, re.IGNORECASE)
            if m:
                degrees.append({
                    "degree": m.group(1),
                    "field": m.group(2).strip()[:60] if len(m.groups()) > 1 else "",
                    "raw": line.strip()[:150],
                })
                break
        else:
            # Look for institution names
            if any(kw in line.lower() for kw in ["university", "college", "institute", "school"]):
                if not any(d.get("raw") == line.strip()[:150] for d in degrees):
                    degrees.append({"degree": "", "field": "", "raw": line.strip()[:150]})
    
    return degrees[:5]


# ── Experience Parsing ────────────────────────────────────────────────────────

def parse_experience(exp_text: str) -> List[Dict]:
    """Parse work experience entries."""
    if not exp_text.strip():
        return []
    
    # Split on company/role patterns (date patterns as delimiters)
    entries = []
    current = []
    date_pattern = re.compile(
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}'
        r'|\b\d{4}\s*[-–]\s*(?:\d{4}|Present|Current)\b',
        re.IGNORECASE
    )
    
    lines = exp_text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # A new entry often starts with a date range or company name
        if date_pattern.search(stripped) or (len(stripped) < 60 and stripped[0].isupper() and not stripped.endswith('.')):
            if current:
                entries.append('\n'.join(current))
                current = []
        current.append(stripped)
    
    if current:
        entries.append('\n'.join(current))
    
    results = []
    for entry in entries[:8]:
        lines_entry = [l for l in entry.split('\n') if l.strip()]
        if not lines_entry:
            continue
        
        company = lines_entry[0] if lines_entry else ""
        role = lines_entry[1] if len(lines_entry) > 1 else ""
        description = '\n'.join(lines_entry[2:]) if len(lines_entry) > 2 else ""
        
        results.append({
            "company": company[:100],
            "role": role[:100],
            "duration": _extract_duration(entry),
            "description": description[:500],
            "skills_used": extract_skills_from_text(entry),
        })
    
    return results


def _extract_duration(text: str) -> str:
    date_pattern = re.compile(
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}'
        r'\s*[-–]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*(?:\d{4}|Present|Current)'
        r'|\b\d{4}\s*[-–]\s*(?:\d{4}|Present|Current)\b',
        re.IGNORECASE
    )
    m = date_pattern.search(text)
    return m.group(0) if m else ""


# ── Project Parsing ───────────────────────────────────────────────────────────

def parse_projects(proj_text: str) -> List[Dict]:
    """Parse project entries from resume."""
    if not proj_text.strip():
        return []
    
    projects = []
    lines = proj_text.split('\n')
    current_proj = None
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Project name: short, usually capitalized, doesn't start with bullet
        is_name_line = (
            len(stripped) < 80
            and stripped[0].isupper()
            and not stripped.startswith(('•', '-', '*', '–'))
            and not re.match(r'^[\d.]+\s', stripped)
        )
        
        if is_name_line and current_proj:
            projects.append(current_proj)
            current_proj = None
        
        if is_name_line:
            # Extract technologies mentioned in parentheses/brackets
            tech_match = re.search(r'[\(\[](.*?)[\)\]]', stripped)
            techs = []
            if tech_match:
                raw_techs = tech_match.group(1)
                techs = [t.strip() for t in re.split(r'[,|/]', raw_techs) if t.strip()]
            current_proj = {
                "name": stripped[:100],
                "description": "",
                "technologies": techs,
            }
        elif current_proj:
            current_proj["description"] += " " + stripped
            # Also extract inline tech mentions
            more_techs = extract_skills_from_text(stripped)
            current_proj["technologies"] = list(set(current_proj["technologies"] + more_techs))
    
    if current_proj:
        projects.append(current_proj)
    
    for proj in projects:
        proj["description"] = proj["description"].strip()[:400]
    
    return projects[:8]


# ── ATS Compatibility ─────────────────────────────────────────────────────────

def analyze_ats_compatibility(text: str, ats_flags: Dict, sections: Dict) -> Dict:
    """Score ATS compatibility and identify issues."""
    issues = []
    score = 100
    
    contact = extract_contact_info(text)
    
    # Contact completeness
    if not contact["email"]:
        issues.append({"type": "error", "msg": "Email address not detected — critical for ATS parsing"})
        score -= 15
    if not contact["phone"]:
        issues.append({"type": "warning", "msg": "Phone number missing"})
        score -= 5
    if not contact["linkedin"]:
        issues.append({"type": "info", "msg": "LinkedIn URL not found — add for recruiter visibility"})
        score -= 3
    
    # Required sections
    required = ["experience", "education", "skills"]
    for sec in required:
        if not sections.get(sec, "").strip():
            issues.append({"type": "error", "msg": f"'{sec.capitalize()}' section not detected"})
            score -= 10
    
    # Recommended sections
    recommended = ["summary", "projects", "certifications"]
    for sec in recommended:
        if not sections.get(sec, "").strip():
            issues.append({"type": "info", "msg": f"Consider adding a '{sec.capitalize()}' section"})
            score -= 2
    
    # Formatting issues
    if ats_flags.get("has_images"):
        issues.append({"type": "error", "msg": "Images/icons detected — ATS cannot parse them"})
        score -= 10
    
    if ats_flags.get("has_tables"):
        issues.append({"type": "warning", "msg": "Tables detected — may cause ATS parsing failures"})
        score -= 8
    
    if ats_flags.get("multi_column"):
        issues.append({"type": "warning", "msg": "Multi-column layout detected — single column preferred for ATS"})
        score -= 8
    
    if ats_flags.get("font_count", 0) > 5:
        issues.append({"type": "info", "msg": "Many font variations detected — simplify for ATS readability"})
        score -= 3
    
    # Length check
    word_count = len(text.split())
    if word_count < 200:
        issues.append({"type": "warning", "msg": "Resume appears very short — add more detail"})
        score -= 10
    elif word_count > 1200:
        issues.append({"type": "info", "msg": "Resume is very long — consider condensing to 1-2 pages"})
        score -= 3
    
    return {
        "score": max(score, 0),
        "issues": issues,
        "contact_completeness": {
            "email": bool(contact["email"]),
            "phone": bool(contact["phone"]),
            "linkedin": bool(contact["linkedin"]),
            "github": bool(contact["github"]),
        },
        "sections_found": [k for k, v in sections.items() if v.strip() and k != "other"],
    }


# ── Master Parse Function ─────────────────────────────────────────────────────

def parse_resume(file_path: str) -> Dict:
    """Full resume parsing pipeline. Returns structured JSON."""
    text, ats_flags = extract_text(file_path)
    sections = split_into_sections(text)
    
    contact = extract_contact_info(text)
    education = parse_education(sections.get("education", ""))
    experience = parse_experience(sections.get("experience", ""))
    projects = parse_projects(sections.get("projects", ""))
    
    # Skills from dedicated section + inferred from experience/projects
    skill_text = (
        sections.get("skills", "")
        + " " + sections.get("experience", "")
        + " " + sections.get("projects", "")
    )
    skills = extract_skills_from_text(skill_text)
    
    certifications = []
    cert_text = sections.get("certifications", "")
    if cert_text:
        certifications = [l.strip() for l in cert_text.split('\n') if l.strip() and len(l.strip()) > 5][:10]
    
    achievements = []
    ach_text = sections.get("achievements", "")
    if ach_text:
        achievements = [l.strip() for l in ach_text.split('\n') if l.strip() and len(l.strip()) > 10][:8]
    
    ats_analysis = analyze_ats_compatibility(text, ats_flags, sections)
    
    return {
        "contact": contact,
        "education": education,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "certifications": certifications,
        "achievements": achievements,
        "ats_compatibility": ats_analysis,
        "raw_text": text[:3000],  # truncated for payload
        "sections_detected": list(sections.keys()),
    }
